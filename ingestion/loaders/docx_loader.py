"""Word 文档 (.docx/.doc) 加载器模块。

解析 Word 文档的段落和表格内容，转换为 LangChain Document 对象列表。
"""

from __future__ import annotations

import os

from langchain_core.documents import Document

from utils.logger_handler import logger


def docx_loader(filepath: str) -> list[Document]:
    """真正落地 .docx / .doc 解析（不再 logger.warning + return []）。

    使用 python-docx 库；解析失败则退回纯文本读取，再失败 return []。
    返回的每个 Document 的 page_content 按 段落+表格单元格 顺序拼接；
    metadata 写入：source=filepath, file_type="docx", file_name=basename(filepath)。

    Args:
        filepath: Word 文档的绝对或相对路径。

    Returns:
        解析得到的 Document 列表。解析失败或无内容时返回空列表。
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        logger.warning("[docx_loader] 缺少 python-docx 依赖，请执行: pip install python-docx")
        return []

    try:
        doc = DocxDocument(filepath)
    except Exception as e:
        logger.warning(f"[docx_loader] python-docx 打开失败，尝试纯文本读取: {e}")
        try:
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            if not text.strip():
                logger.warning(f"[docx_loader] 纯文本读取也没有内容: {filepath}")
                return []
            return [
                Document(
                    page_content=text,
                    metadata={
                        "source": filepath,
                        "file_type": "docx",
                        "file_name": os.path.basename(filepath),
                    },
                )
            ]
        except Exception as e2:
            logger.warning(f"[docx_loader] 纯文本读取也失败: {e2}")
            return []

    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text + "\n")

    for table in doc.tables:
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_cells.append(cell_text)
            if row_cells:
                parts.append(" | ".join(row_cells) + "\n")

    full_text = "\n".join(parts)

    if not full_text.strip():
        logger.warning(f"[docx_loader] docx 没提取到文字: {filepath}")
        return []

    return [
        Document(
            page_content=full_text,
            metadata={
                "source": filepath,
                "file_type": "docx",
                "file_name": os.path.basename(filepath),
            },
        )
    ]
